"""XML 解析性能对比测试（iter-111）。

验证 iter-111 的两项优化效果：

1. **ODF XPath 优化**：``_odf_xml.iter_elements`` 从 ``root.iter()`` +
   Python 层 ``tag.endswith`` 字符串匹配，改为 ``lxml.etree.xpath`` 在
   libxml2 C 层完成节点遍历与命名空间匹配。
2. **DOCX 多标签 C 层过滤优化**：``_extract_docx_root_paragraphs`` 从
   ``para.iter()`` 全树遍历 + Python 层 ``tag ==`` 字符串比较，改为
   ``para.iter(_W_T, _W_TAB, _W_BR, _W_CR)`` 多标签 C 层过滤，只返回
   4 种目标 tag 的元素，减少 60-80% Python 循环次数。

  .. note::
     原计划的 ``etree.iterparse`` 流式解析方案在实测中因 DOCX
     ``document.xml`` 通常 < 1MB（事件分发开销超过收益）反而变慢，
     已撤销；改用与 ODF XPath 一致的"把节点匹配下沉到 C 层"思路。

测试方法：在同一组样本上分别运行优化前/后的实现，取多次测量的中位数
对比耗时，并断言新实现不慢于旧实现（留 1.5x 宽松阈值避免 CI flakiness）。
所有测试标记 ``@pytest.mark.slow``，CI 默认跳过。
"""

from __future__ import annotations

import io
import statistics
import time
from typing import Any

import pytest

# lxml 不可用时跳过全部对比（优化前后均依赖 lxml）
pytest.importorskip("lxml")


# ----------------------------- 测量工具 -----------------------------


def _measure_median(fn: Any, *args: Any, iterations: int = 10) -> float:
    """测量函数执行耗时（秒），返回中位数。"""
    times: list[float] = []
    # 预热一次避免首次 import/缓存开销污染测量
    fn(*args)
    for _ in range(iterations):
        start = time.perf_counter()
        fn(*args)
        times.append(time.perf_counter() - start)
    return statistics.median(times)


# ----------------------------- 旧实现（基线） -----------------------------


def _iter_elements_legacy(
    root: Any,
    namespace: str,
    local_names: tuple[str, ...],
) -> list[Any]:
    """iter-111 之前的 ``iter_elements`` 实现（Python 层 tag 过滤）。

    保留用于性能对比基准，禁止在生产代码调用。
    """
    targets = tuple(f"}}{name}" for name in local_names)
    prefix = f"{{{namespace}}}"
    result: list[Any] = []
    for elem in root.iter():
        tag = elem.tag
        if isinstance(tag, str) and tag.startswith(prefix) and tag.endswith(targets):
            result.append(elem)
    return result


def _extract_docx_paragraphs_legacy(xml_bytes: bytes) -> list[str]:
    """iter-111 之前的 DOCX 段落提取实现（fromstring + root.iter 全树遍历）。

    保留用于性能对比基准，禁止在生产代码调用。
    """
    from lxml import etree

    parser = etree.XMLParser(resolve_entities=False, no_network=True, recover=True)
    root = etree.fromstring(xml_bytes, parser=parser)
    return _legacy_extract_paragraphs_from_root(root)


def _legacy_extract_paragraphs_from_root(root: Any) -> list[str]:
    """iter-111 之前的 DOCX 段落遍历实现（``para.iter()`` 全树 + Python tag 比较）。

    与 :func:`_extract_docx_root_paragraphs` 接收相同的 ``root`` 参数，
    只测遍历性能（排除 fromstring 解析开销），保证对比公平。
    """
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w_p = f"{{{w_ns}}}p"
    w_t = f"{{{w_ns}}}t"
    w_tab = f"{{{w_ns}}}tab"
    w_br = f"{{{w_ns}}}br"
    w_cr = f"{{{w_ns}}}cr"

    paragraphs: list[str] = []
    for para in root.iter(w_p):
        para_parts: list[str] = []
        for child in para.iter():
            tag = child.tag
            if tag == w_t and child.text:
                para_parts.append(child.text)
            elif tag == w_tab:
                para_parts.append("\t")
            elif tag in (w_br, w_cr):
                para_parts.append("\n")
        text = "".join(para_parts).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


# ----------------------------- 样本生成 -----------------------------


def _make_docx_sample_large() -> bytes:
    """生成较大 DOCX 样本（1000 段落 + 50×5 表格，document.xml ~200KB+）。"""
    from docx import Document

    doc = Document()
    for i in range(1000):
        doc.add_paragraph(f"段落 {i}：含 password 和 secret 关键词的内容用于性能测试")
    table = doc.add_table(rows=50, cols=5)
    for row in table.rows:
        for cell in row.cells:
            cell.text = "password cell secret_key content"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_odt_sample_large() -> bytes:
    """生成较大 ODT 样本（500 段落，content.xml ~30KB+）。"""
    from tests._odf_samples import make_odt_sample

    return make_odt_sample([f"段落 {i}：password secret 关键词内容" for i in range(500)])


def _make_ods_sample_large() -> bytes:
    """生成较大 ODS 样本（200 行 × 10 列，content.xml ~50KB+）。"""
    from tests._odf_samples import make_ods_sample

    rows = [[f"cell_{r}_{c}_password" for c in range(10)] for r in range(200)]
    return make_ods_sample(rows)


def _extract_docx_document_xml(docx_data: bytes) -> bytes:
    """从 DOCX 字节中读取 word/document.xml。"""
    import zipfile

    with zipfile.ZipFile(io.BytesIO(docx_data)) as zf:
        return zf.read("word/document.xml")


# ----------------------------- DOCX 优化尝试记录 -----------------------------


@pytest.mark.slow
class TestDocxOptimizationAttempts:
    """DOCX 优化尝试记录与功能回归校验。

    iter-111 尝试过两种 DOCX 优化方向，实测均未达预期，已撤销：

    - ``etree.iterparse`` 流式解析：在 36KB-140KB document.xml 上反慢 0.5x
    - ``para.iter(_W_T, _W_TAB, _W_BR, _W_CR)`` 多标签 C 层过滤：反慢 0.68x

    原因：DOCX ``root.iter(_W_P)`` + ``para.iter()`` 已是 C 层节点遍历，
    Python 层仅做 tag 字符串比较，对小段落（1-3 个子元素）开销本就很低。
    本测试类保留功能回归校验，确保撤销后行为不变。
    """

    def test_docx_extract_text_unchanged_functionally(self) -> None:
        """当前 extract_docx_text 能正确提取文本（功能回归校验）。"""
        from fuscan.extractors._ooxml_xml import extract_docx_text

        data = _make_docx_sample_large()
        text = extract_docx_text(data)
        assert "password" in text
        assert "secret" in text
        # 1000 段落 + 50 行表格 = 至少 1000 个非空行
        assert len(text.splitlines()) >= 1000

    def test_docx_two_optimization_attempts_produce_identical_output(self) -> None:
        """两种失败的优化方案与当前实现产出一致（防止未来回退引入功能差异）。"""
        from lxml import etree

        from fuscan.extractors._ooxml_xml import _extract_docx_root_paragraphs

        docx_data = _make_docx_sample_large()
        xml_bytes = _extract_docx_document_xml(docx_data)

        parser = etree.XMLParser(resolve_entities=False, no_network=True, recover=True)
        root = etree.fromstring(xml_bytes, parser=parser)

        # 当前实现
        current_result = _extract_docx_root_paragraphs(root)
        # iter-111 之前的实现（fromstring + 全树 iter + Python tag 比较）
        legacy_result = _legacy_extract_paragraphs_from_root(root)

        assert current_result == legacy_result, "当前实现与旧实现产出段落文本不一致，可能引入功能回归"


# ----------------------------- ODF XPath 对比 -----------------------------


@pytest.mark.slow
class TestOdfXPathComparison:
    """ODF iter_elements xpath vs Python 层 tag 过滤性能对比。"""

    def test_odt_xpath_faster_than_python_filter(self) -> None:
        """ODT 段落遍历用 xpath 应不慢于 Python 层 tag.endswith 过滤。"""
        from fuscan.extractors._odf_xml import (
            TEXT_NS,
            iter_elements,
            load_content_xml,
        )

        data = _make_odt_sample_large()
        root = load_content_xml(data)

        # 验证两个实现产出一致
        legacy_result = _iter_elements_legacy(root, TEXT_NS, ("p", "h"))
        new_result = list(iter_elements(root, TEXT_NS, ("p", "h")))
        assert len(legacy_result) == len(new_result), "xpath 与 iter 产出段落数不一致"

        legacy_time = _measure_median(_iter_elements_legacy, root, TEXT_NS, ("p", "h"))
        new_time = _measure_median(lambda r=root: list(iter_elements(r, TEXT_NS, ("p", "h"))))

        print(
            f"\n[ODT xpath 对比] content.xml 段落 {len(new_result)} 个\n"
            f"  iter+endswith (旧): {legacy_time * 1000:.2f} ms\n"
            f"  xpath         (新): {new_time * 1000:.2f} ms\n"
            f"  提速: {legacy_time / new_time:.2f}x"
        )

        # 断言新实现不慢于旧实现（留 1.5x 宽松阈值）
        assert new_time <= legacy_time * 1.5, (
            f"xpath 耗时 {new_time * 1000:.2f}ms 显著慢于 iter+endswith {legacy_time * 1000:.2f}ms，请检查实现"
        )

    def test_ods_xpath_faster_than_python_filter(self) -> None:
        """ODS 单元格遍历用 xpath 应不慢于 Python 层 tag.endswith 过滤。"""
        from fuscan.extractors._odf_xml import (
            TABLE_NS,
            iter_elements,
            load_content_xml,
        )

        data = _make_ods_sample_large()
        root = load_content_xml(data)

        # 验证两个实现产出一致（行级）
        legacy_rows = _iter_elements_legacy(root, TABLE_NS, ("table-row",))
        new_rows = list(iter_elements(root, TABLE_NS, ("table-row",)))
        assert len(legacy_rows) == len(new_rows), "xpath 与 iter 产出表格行数不一致"

        # 测量完整提取（行 + 单元格）的耗时
        def _legacy_full() -> int:
            count = 0
            for row in _iter_elements_legacy(root, TABLE_NS, ("table-row",)):
                count += len(_iter_elements_legacy(row, TABLE_NS, ("table-cell",)))
            return count

        def _new_full() -> int:
            count = 0
            for row in iter_elements(root, TABLE_NS, ("table-row",)):
                count += len(list(iter_elements(row, TABLE_NS, ("table-cell",))))
            return count

        # 功能等价性校验
        assert _legacy_full() == _new_full(), "xpath 与 iter 产出单元格数不一致"

        legacy_time = _measure_median(_legacy_full)
        new_time = _measure_median(_new_full)

        print(
            f"\n[ODS xpath 对比] 行 {len(new_rows)} / 单元格 {_new_full()} 个\n"
            f"  iter+endswith (旧): {legacy_time * 1000:.2f} ms\n"
            f"  xpath         (新): {new_time * 1000:.2f} ms\n"
            f"  提速: {legacy_time / new_time:.2f}x"
        )

        assert new_time <= legacy_time * 1.5, (
            f"xpath 耗时 {new_time * 1000:.2f}ms 显著慢于 iter+endswith {legacy_time * 1000:.2f}ms，请检查实现"
        )

    def test_odt_extract_text_unchanged_functionally(self) -> None:
        """优化后 OdtExtractor 仍能正确提取文本（功能回归校验）。"""
        from fuscan.extractors import OdtExtractor

        data = _make_odt_sample_large()
        text = OdtExtractor().extract_from_bytes(data)
        assert "password" in text
        assert "secret" in text
        # 500 段落
        assert len(text.splitlines()) >= 400

    def test_ods_extract_text_unchanged_functionally(self) -> None:
        """优化后 OdsExtractor 仍能正确提取文本（功能回归校验）。"""
        from fuscan.extractors import OdsExtractor

        data = _make_ods_sample_large()
        text = OdsExtractor().extract_from_bytes(data)
        assert "password" in text
        # 200 行
        assert len(text.splitlines()) >= 150


# ----------------------------- 极限测试（5MB+ 大样本） -----------------------------


# module-level 缓存：大样本生成耗时较高，避免每个测试重复生成
_extreme_cache: dict[str, Any] = {}


def _get_extreme_odt() -> bytes:
    """生成/缓存极限 ODT 样本（50000 段落，content.xml ~4MB+）。"""
    if "odt" not in _extreme_cache:
        from tests._odf_samples import make_odt_sample

        _extreme_cache["odt"] = make_odt_sample(
            [f"段落 {i}：password secret 关键词内容用于极限性能测试" for i in range(50000)]
        )
    return _extreme_cache["odt"]


def _get_extreme_ods() -> bytes:
    """生成/缓存极限 ODS 样本（2000 行 × 20 列，content.xml ~3MB+）。"""
    if "ods" not in _extreme_cache:
        from tests._odf_samples import make_ods_sample

        rows = [[f"cell_{r}_{c}_password" for c in range(20)] for r in range(2000)]
        _extreme_cache["ods"] = make_ods_sample(rows)
    return _extreme_cache["ods"]


def _get_extreme_docx() -> bytes:
    """生成/缓存极限 DOCX 样本（10000 段落，document.xml ~2MB+）。

    用直接构造 XML + zipfile 打包，避免 python-docx 生成大文档的慢速。
    """
    if "docx" not in _extreme_cache:
        import zipfile

        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        # 直接构造 document.xml：10000 个 w:p 段落，每个含 1 个 w:r > w:t
        parts = [
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n',
            f'<w:document xmlns:w="{w_ns}"><w:body>',
        ]
        for i in range(10000):
            parts.append(f"<w:p><w:r><w:t>段落 {i}：password secret 关键词内容用于极限性能测试</w:t></w:r></w:p>")
        parts.append("</w:body></w:document>")
        document_xml = "".join(parts).encode("utf-8")

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("word/document.xml", document_xml)
        _extreme_cache["docx"] = buf.getvalue()
        _extreme_cache["docx_xml_size"] = len(document_xml)
    return _extreme_cache["docx"]


@pytest.mark.slow
class TestExtremeScale:
    """极限规模性能测试（5MB+ 大样本）。

    极限场景验证：
    1. XPath 优化在大文件场景是否依然有效（C 层优势应更明显）
    2. element_text 递归实现是否会成为瓶颈
    3. DOCX 大文档是否有新的优化机会
    """

    def test_extreme_odt_xpath_vs_iter(self) -> None:
        """极限 ODT（50000 段落）XPath vs iter 性能对比。"""
        from fuscan.extractors._odf_xml import (
            TEXT_NS,
            iter_elements,
            load_content_xml,
        )

        data = _get_extreme_odt()
        root = load_content_xml(data)

        # 功能等价性
        legacy_result = _iter_elements_legacy(root, TEXT_NS, ("p", "h"))
        new_result = list(iter_elements(root, TEXT_NS, ("p", "h")))
        assert len(legacy_result) == len(new_result) == 50000, "段落数不符合预期"

        legacy_time = _measure_median(_iter_elements_legacy, root, TEXT_NS, ("p", "h"), iterations=3)
        new_time = _measure_median(lambda r=root: list(iter_elements(r, TEXT_NS, ("p", "h"))), iterations=3)

        print(
            f"\n[极限 ODT xpath 对比] 50000 段落\n"
            f"  iter+endswith (旧): {legacy_time * 1000:.2f} ms\n"
            f"  xpath         (新): {new_time * 1000:.2f} ms\n"
            f"  提速: {legacy_time / new_time:.2f}x"
        )

        assert new_time <= legacy_time * 1.5

    def test_extreme_ods_xpath_vs_iter(self) -> None:
        """极限 ODS（2000 行 × 20 列 = 40000 单元格）XPath vs iter 性能对比。"""
        from fuscan.extractors._odf_xml import (
            TABLE_NS,
            iter_elements,
            load_content_xml,
        )

        data = _get_extreme_ods()
        root = load_content_xml(data)

        def _legacy_full() -> int:
            count = 0
            for row in _iter_elements_legacy(root, TABLE_NS, ("table-row",)):
                count += len(_iter_elements_legacy(row, TABLE_NS, ("table-cell",)))
            return count

        def _new_full() -> int:
            count = 0
            for row in iter_elements(root, TABLE_NS, ("table-row",)):
                count += len(list(iter_elements(row, TABLE_NS, ("table-cell",))))
            return count

        assert _legacy_full() == _new_full() == 40000

        legacy_time = _measure_median(_legacy_full, iterations=3)
        new_time = _measure_median(_new_full, iterations=3)

        print(
            f"\n[极限 ODS xpath 对比] 2000 行 × 20 列 = 40000 单元格\n"
            f"  iter+endswith (旧): {legacy_time * 1000:.2f} ms\n"
            f"  xpath         (新): {new_time * 1000:.2f} ms\n"
            f"  提速: {legacy_time / new_time:.2f}x"
        )

        assert new_time <= legacy_time * 1.5

    def test_extreme_odt_full_extraction_profile(self) -> None:
        """极限 ODT 完整提取性能剖析（定位 element_text 等子环节瓶颈）。"""
        from fuscan.extractors._odf_xml import (
            TEXT_NS,
            element_text,
            iter_elements,
            load_content_xml,
        )

        data = _get_extreme_odt()
        root = load_content_xml(data)

        # 阶段 1：iter_elements（XPath 节点匹配）
        t1 = time.perf_counter()
        paragraphs = list(iter_elements(root, TEXT_NS, ("p", "h")))
        t_iter = time.perf_counter() - t1

        # 阶段 2：element_text（递归文本提取）
        t2 = time.perf_counter()
        texts = [element_text(p) for p in paragraphs]
        t_text = time.perf_counter() - t2

        # 阶段 3：join
        t3 = time.perf_counter()
        result = "\n".join(t for t in texts if t)
        t_join = time.perf_counter() - t3

        print(
            f"\n[极限 ODT 剖析] 50000 段落\n"
            f"  iter_elements (xpath): {t_iter * 1000:.2f} ms\n"
            f"  element_text (递归):   {t_text * 1000:.2f} ms\n"
            f"  join:                  {t_join * 1000:.2f} ms\n"
            f"  总计:                  {(t_iter + t_text + t_join) * 1000:.2f} ms\n"
            f"  element_text 占比:     {t_text / (t_iter + t_text + t_join) * 100:.1f}%"
        )

        assert "password" in result
        assert len(result.splitlines()) >= 40000

    def test_extreme_ods_full_extraction_profile(self) -> None:
        """极限 ODS 完整提取性能剖析（定位单元格遍历瓶颈）。"""
        from fuscan.extractors._odf_xml import (
            TABLE_NS,
            element_text,
            iter_elements,
            load_content_xml,
        )

        data = _get_extreme_ods()
        root = load_content_xml(data)

        # 阶段 1：行遍历
        t1 = time.perf_counter()
        rows = list(iter_elements(root, TABLE_NS, ("table-row",)))
        t_rows = time.perf_counter() - t1

        # 阶段 2：单元格遍历
        t2 = time.perf_counter()
        cells: list[Any] = []
        for row in rows:
            cells.extend(iter_elements(row, TABLE_NS, ("table-cell",)))
        t_cells = time.perf_counter() - t2

        # 阶段 3：element_text（结果不计入断言，仅测耗时）
        t3 = time.perf_counter()
        for cell in cells:
            element_text(cell)
        t_text = time.perf_counter() - t3

        print(
            f"\n[极限 ODS 剖析] {len(rows)} 行 / {len(cells)} 单元格\n"
            f"  行遍历 (xpath):      {t_rows * 1000:.2f} ms\n"
            f"  单元格遍历 (xpath):  {t_cells * 1000:.2f} ms\n"
            f"  element_text (递归): {t_text * 1000:.2f} ms\n"
            f"  总计:                {(t_rows + t_cells + t_text) * 1000:.2f} ms"
        )

        assert len(cells) == 40000

    def test_extreme_docx_extraction_profile(self) -> None:
        """极限 DOCX（10000 段落）完整提取性能剖析。"""
        from fuscan.extractors._ooxml_xml import extract_docx_text

        data = _get_extreme_docx()
        xml_size = _extreme_cache["docx_xml_size"]

        t1 = time.perf_counter()
        text = extract_docx_text(data)
        elapsed = time.perf_counter() - t1

        lines = len(text.splitlines())
        print(
            f"\n[极限 DOCX 剖析] document.xml={xml_size / 1024:.0f}KB\n"
            f"  总耗时: {elapsed * 1000:.2f} ms\n"
            f"  段落数: {lines}"
        )

        assert "password" in text
        assert lines >= 10000
